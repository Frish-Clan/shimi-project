#!/usr/bin/env python3
"""Generate Hebrew instability index report as DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup (A4, RTL) ─────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2)
section.bottom_margin = Cm(2)

# Enable RTL for the document
doc.core_properties.language = 'he-IL'

# ── Styles ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def heading1(text, color=(31, 73, 125)):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Arial'
    return p

def heading2(text, color=(0, 70, 127)):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Arial'
    return p

def heading3(text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    return p

def body(text, bold=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def divider():
    p = doc.add_paragraph('─' * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)

def add_table(headers, rows, col_widths=None, header_color=(31, 73, 125)):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '%02X%02X%02X' % header_color)
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = 'F0F4FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            # Alternate row shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg)
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

    # Set column widths if provided
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('מדד אי-היציבות — Shimi Dashboard')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 73, 125)
run.font.name = 'Arial'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('דוח מתודולוגי — מקורות נתונים ומחשוב ניקוד')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(89, 89, 89)
run2.font.name = 'Arial'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(f'תאריך: {datetime.date.today().strftime("%d/%m/%Y")}')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(120, 120, 120)
run3.font.name = 'Arial'

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

heading1('1. סקירה כללית')
divider()
body(
    'מדד אי-היציבות מחשב ציון מורכב בסולם 0–100 לכל מדינה, '
    'המורכב ממספר תת-מדדים ואותות חיים ממקורות מידע מגוונים. '
    'הדוח הבא מפרט את מקורות הנתונים, שיטת החישוב, ואת ההשוואה '
    'לאתר worldmonitor.app.'
)

body('הנוסחה הבסיסית:')
p = doc.add_paragraph()
set_rtl(p)
run = p.add_run('composite = max(floor, min(100, baseline×0.4 + eventScore×0.6 + boosts))')
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.bold = True

body('כאשר eventScore = U×0.25 + C×0.30 + S×0.20 + I×0.25')

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SUB-INDICES
# ════════════════════════════════════════════════════════════════════════════

heading1('2. תת-מדדים (U / C / S / I)')
divider()

add_table(
    ['תת-מדד', 'שם', 'משקל', 'המדד שלנו', 'worldmonitor'],
    [
        ['U', 'Unrest — אי-שקט', '25%', 'אירועי מחאה מ-Bootstrap (ciiContribution)', 'ACLED Live + Liveuamap בזמן אמת'],
        ['C', 'Conflict — עימות', '30%', 'ACLED מצטבר מ-Bootstrap (geoConvergence)', 'ACLED Live: קרבות, פיצוצים, קורבנות'],
        ['S', 'Security — ביטחון', '20%', 'GPS Jamming מ-/api/gpsjam (militaryActivity)', 'ADS-B טיסות + AIS כלי שייט + GPS חיות'],
        ['I', 'Information — מידע', '25%', 'נפח חדשות מ-Bootstrap (newsActivity)', 'אשכולות חדשות חיות + מהירות מקורות'],
    ],
    col_widths=[1.5, 3.5, 1.5, 5, 5]
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DATA SOURCES
# ════════════════════════════════════════════════════════════════════════════

heading1('3. מקורות נתונים — פירוט')
divider()

heading2('3.1 מקורות פעילים במדד שלנו ✅')

add_table(
    ['מקור', 'נתון', 'תרומה לציון', 'עדכון'],
    [
        ['worldmonitor /api/bootstrap', 'CII scores, UCDP, unrest, wildfires', 'ציון בסיס + תת-מדדים', '10 דקות'],
        ['worldmonitor /api/gpsjam', 'היקס H3 של הפרעות GPS', 'עד +35 לתת-מדד S', '15 דקות'],
        ['worldmonitor /api/bootstrap → cyberThreats', 'איומי סייבר לפי מדינה', 'עד +12 לפי חומרה', '10 דקות'],
        ['worldmonitor /api/bootstrap → outages', 'הפסקות אינטרנט לפי מדינה', 'עד +50 (הפסקה לאומית)', '10 דקות'],
        ['worldmonitor /api/bootstrap → temporalAnomalies', 'חריגות זמניות', 'עד +6', '10 דקות'],
        ['US State Dept RSS', 'רמות אזהרת מסע (Level 1–4)', 'עד +15 + רצפה 60', '2 שעות'],
        ['UK FCDO Atom Feed', 'אזהרות מסע בריטיות', '+3 תוספת אם מרובי-מקורות', '2 שעות'],
        ['NASA USGS Earthquakes', 'רעידות אדמה', 'אות אקלים (signals)', '15 דקות'],
        ['World Bank (PV/RL/CC.EST)', 'ממשל ושלטון חוק', 'גיבוי כשאין CII', 'שנתי'],
        ['GDELT API', 'ציר זמן עימותים 30 יום', 'גיבוי לחדשות', '2 שעות'],
    ],
    col_widths=[4.5, 4, 3.5, 2]
)

heading2('3.2 מקורות חסרים — קיימים ב-worldmonitor בלבד ❌')

add_table(
    ['מקור', 'תרומה לציון', 'סיבת החסר'],
    [
        ['Oref — פיקוד העורף (ישראל)', 'עד +50 לישראל בזמן התראה', 'API פנימי — לא נגיש'],
        ['ACLED Live Events', 'עמוד שדרה של תת-מדד C', 'מנוי בתשלום ($)'],
        ['ADS-B טיסות צבאיות', 'תת-מדד S בזמן אמת', 'מנוי בתשלום ($)'],
        ['AIS כלי שייט', 'תת-מדד S ימי', 'מנוי בתשלום ($)'],
        ['Liveuamap Events', 'אירועי פגיעה ממוקמים גיאוגרפית', 'API פנימי — לא נגיש'],
        ['מנוע אשכולות חדשות פנימי', 'תת-מדד I — velocity', 'עיבוד פנימי בלבד'],
        ['Security Advisories AU/NZ', '+5 אם 3+ מקורות', 'CORS חוסם גישה ישירה'],
        ['HAPI (Humanitarian API)', 'נתוני עקורים ומשבר הומניטרי', 'ניסיון אך לא מחובר'],
    ],
    col_widths=[4.5, 4, 5.5],
    header_color=(192, 0, 0)
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — EXAMPLE ISRAEL
# ════════════════════════════════════════════════════════════════════════════

heading1('4. דוגמת חישוב — ישראל (IL)')
divider()

add_table(
    ['רכיב', 'המדד שלנו', 'worldmonitor', 'הסבר'],
    [
        ['ציון מורכב סופי', '~51', '~70', 'פער של ~19 נקודות'],
        ['CII Baseline (שרת)', '45', '45', 'זהה — אותו API endpoint'],
        ['U — Unrest', '0', '~29', 'ACLED: אין לנו מחאות חיות'],
        ['C — Conflict', '50', '~42', 'ACLED מצטבר vs. עדכונים חיים'],
        ['S — Security', '24', '~22', 'GPS jamming דומה — טיסות חסרות'],
        ['I — Information', '0', '~54', 'חסר מנוע אשכולות חדשות פנימי'],
        ['Cyber Threats Boost', '+6', '+6', '2 איומי סייבר קריטיים — זהה'],
        ['Advisory Boost', '0', '+10', 'ישראל אינה ב-RSS הראשי (עמוד נפרד)'],
        ['Oref Alerts Boost', '0', '+15 עד +40', 'אין גישה ל-API פיקוד העורף'],
        ['ACLED Live Boost', '0', '+8 עד +15', 'אין מנוי ACLED'],
    ],
    col_widths=[4, 2.5, 2.5, 5]
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EXAMPLE UKRAINE
# ════════════════════════════════════════════════════════════════════════════

heading1('5. דוגמת חישוב — אוקראינה (UA)')
divider()

add_table(
    ['רכיב', 'המדד שלנו', 'worldmonitor', 'הסבר'],
    [
        ['ציון מורכב סופי', '~60', '~75', 'פער של ~15 נקודות'],
        ['CII Baseline (שרת)', '50', '50', 'זהה'],
        ['Advisory Floor (Do Not Travel)', '60', '60', 'US State Dept Level 4 — זהה'],
        ['GPS Jamming Boost', '+8 (S sub-index)', '+10', 'כמות היקסים דומה'],
        ['Advisory Boost', '+10', '+10', 'Level 3 Reconsider Travel'],
        ['ACLED + Oref', '0', '+15', 'קרבות חיים, פגיעות — אין לנו'],
    ],
    col_widths=[4.5, 2.5, 2.5, 4.5]
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — COVERAGE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

heading1('6. סיכום כיסוי מקורות')
divider()

add_table(
    ['קטגוריה', 'כיסוי שלנו', 'כיסוי worldmonitor'],
    [
        ['נתוני ממשל שנתיים (World Bank)', '✅ מלא', '✅ מלא'],
        ['ציון CII בסיסי (שרת)', '✅ מלא', '✅ מלא'],
        ['GPS Jamming (H3 hex)', '✅ מלא', '✅ מלא'],
        ['אזהרות מסע (US+UK)', '✅ חלקי', '✅ מלא (US+UK+AU+NZ)'],
        ['איומי סייבר', '✅ לפי מדינה', '✅ מלא'],
        ['שריפות לוויין FIRMS', '✅ לפי מדינה (radius)', '✅ מלא'],
        ['הפסקות אינטרנט', '✅ לפי מדינה', '✅ מלא'],
        ['נתוני עימות ACLED', '⚠️ מצטבר בלבד', '✅ חיות'],
        ['נתוני UCDP', '✅ חלקי', '✅ מלא'],
        ['התראות טילים Oref (ישראל)', '❌ אין גישה', '✅ בזמן אמת'],
        ['טיסות צבאיות ADS-B', '❌ אין גישה', '✅ בזמן אמת'],
        ['אשכולות חדשות חיות', '❌ אין גישה', '✅ בזמן אמת'],
        ['כלי שייט AIS', '❌ אין גישה', '✅ בזמן אמת'],
    ],
    col_widths=[6, 4, 4]
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — CONCLUSIONS
# ════════════════════════════════════════════════════════════════════════════

heading1('7. מסקנות')
divider()

body(
    'המדד שלנו מכסה כ-70% מהאותות שWorldMonitor משתמש בהם. '
    'הפער העיקרי (~15–25 נקודות) נובע משלושה מקורות קריטיים:',
    bold=False
)

for item in [
    '1. Oref (פיקוד העורף) — רלוונטי בעיקר לישראל, API פנימי של worldmonitor',
    '2. ACLED Live — נתוני עימות בזמן אמת, מצריך מנוי בתשלום',
    '3. מנוע אשכולות חדשות פנימי — מחשב I sub-index, לא נגיש',
]:
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)

doc.add_paragraph()
body(
    'הנתונים שאנחנו מחשבים הם אמינים ומדויקים — הם מבוססים על מקורות פתוחים '
    'ועל ה-API הרשמי של worldmonitor. הציונים משקפים נאמנה את המציאות עבור '
    'רוב המדינות, עם חריגה ידועה למדינות שמצויות בסכסוך פעיל (ישראל, אוקראינה) '
    'שם האותות החיים של worldmonitor מוסיפים משקל משמעותי.',
    bold=False
)

# ── Footer note ──────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Shimi Dashboard • נוצר אוטומטית • {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.name = 'Arial'

# ── Save ─────────────────────────────────────────────────────────────────────
out = '/Users/emanuell/dev/shimi_project/instability_index_report.docx'
doc.save(out)
print(f'Saved: {out}')
